"""
Text extraction service for PDF and EPUB files.
Extracts text content for full-text search indexing.
"""
import io
import hashlib
import asyncio
from pathlib import Path
from typing import Any
from functools import lru_cache
from datetime import datetime, timedelta
import fitz  # PyMuPDF
from ebooklib import epub
from bs4 import BeautifulSoup
from docx import Document
import aiofiles
import json
import os

from app.core.config import settings


class TextExtractionError(Exception):
    """Error during text extraction."""
    pass


class TextCache:
    """Simple file-based cache for extracted text."""

    def __init__(self, cache_dir: str = "/tmp/text_cache"):
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        self.cache_ttl = timedelta(days=7)

    def _get_cache_key(self, file_path: str, file_hash: str) -> str:
        """Generate cache key from file path and hash."""
        combined = f"{file_path}:{file_hash}"
        return hashlib.md5(combined.encode()).hexdigest()

    def _get_cache_path(self, cache_key: str) -> Path:
        """Get cache file path."""
        return self.cache_dir / f"{cache_key}.json"

    async def get(self, file_path: str, file_hash: str) -> dict[str, Any] | None:
        """Get cached extraction result."""
        cache_key = self._get_cache_key(file_path, file_hash)
        cache_path = self._get_cache_path(cache_key)

        if not cache_path.exists():
            return None

        try:
            async with aiofiles.open(cache_path, 'r') as f:
                data = json.loads(await f.read())

            # Check expiration
            cached_at = datetime.fromisoformat(data.get("cached_at", ""))
            if datetime.utcnow() - cached_at > self.cache_ttl:
                cache_path.unlink()
                return None

            return data.get("result")
        except Exception:
            return None

    async def set(self, file_path: str, file_hash: str, result: dict[str, Any]):
        """Cache extraction result."""
        cache_key = self._get_cache_key(file_path, file_hash)
        cache_path = self._get_cache_path(cache_key)

        data = {
            "cached_at": datetime.utcnow().isoformat(),
            "file_path": file_path,
            "file_hash": file_hash,
            "result": result,
        }

        try:
            async with aiofiles.open(cache_path, 'w') as f:
                await f.write(json.dumps(data, ensure_ascii=False))
        except Exception:
            pass  # Caching failures are not critical

    async def invalidate(self, file_path: str, file_hash: str):
        """Invalidate cache entry."""
        cache_key = self._get_cache_key(file_path, file_hash)
        cache_path = self._get_cache_path(cache_key)
        if cache_path.exists():
            cache_path.unlink()


class TextExtractor:
    """
    Service for extracting text content from PDF, EPUB, TXT, and DOCX files.
    Uses PyMuPDF for PDF, ebooklib for EPUB, and python-docx for DOCX.
    """

    # Maximum text length to extract (for search indexing)
    MAX_TEXT_LENGTH = 500_000  # ~500KB of text

    # Maximum pages to process for PDF
    MAX_PDF_PAGES = 500

    def __init__(self):
        self._cache = TextCache()

    def _compute_file_hash(self, content: bytes) -> str:
        """Compute MD5 hash of file content."""
        return hashlib.md5(content).hexdigest()

    async def extract_text(
        self,
        content: bytes,
        content_type: str,
        file_name: str,
        use_cache: bool = True,
    ) -> dict[str, Any]:
        """
        Extract text from file content.

        Args:
            content: File content bytes
            content_type: MIME type of the file
            file_name: Original file name
            use_cache: Whether to use caching

        Returns:
            Dictionary with:
            - text: Extracted text content
            - page_count: Number of pages (if applicable)
            - metadata: Extracted metadata
            - extraction_method: Method used for extraction
        """
        file_hash = self._compute_file_hash(content)

        # Check cache
        if use_cache:
            cached = await self._cache.get(file_name, file_hash)
            if cached:
                return cached

        # Determine file type and extract
        result: dict[str, Any]

        if content_type == "application/pdf" or file_name.lower().endswith(".pdf"):
            result = await self._extract_pdf(content)
        elif content_type == "application/epub+zip" or file_name.lower().endswith(".epub"):
            result = await self._extract_epub(content)
        elif content_type == "text/plain" or file_name.lower().endswith(".txt"):
            result = await self._extract_text_file(content)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.lower().endswith(".docx"):
            result = await self._extract_docx(content)
        else:
            raise TextExtractionError(f"Unsupported file type: {content_type}")

        # Truncate if too long
        if len(result.get("text", "")) > self.MAX_TEXT_LENGTH:
            result["text"] = result["text"][:self.MAX_TEXT_LENGTH]
            result["truncated"] = True

        # Cache result
        if use_cache:
            await self._cache.set(file_name, file_hash, result)

        return result

    async def _extract_pdf(self, content: bytes) -> dict[str, Any]:
        """Extract text from PDF using PyMuPDF."""
        def _do_extraction() -> dict[str, Any]:
            text_parts = []
            metadata = {}

            try:
                doc = fitz.open(stream=content, filetype="pdf")

                # Extract metadata
                meta = doc.metadata
                if meta:
                    metadata = {
                        "title": meta.get("title"),
                        "author": meta.get("author"),
                        "subject": meta.get("subject"),
                        "keywords": meta.get("keywords"),
                        "creator": meta.get("creator"),
                        "producer": meta.get("producer"),
                        "creation_date": meta.get("creationDate"),
                        "mod_date": meta.get("modDate"),
                    }
                    # Clean None values
                    metadata = {k: v for k, v in metadata.items() if v}

                page_count = len(doc)

                # Extract text from each page
                for page_num, page in enumerate(doc):
                    if page_num >= self.MAX_PDF_PAGES:
                        break

                    text = page.get_text("text")
                    if text.strip():
                        text_parts.append(text)

                doc.close()

                return {
                    "text": "\n\n".join(text_parts),
                    "page_count": page_count,
                    "metadata": metadata,
                    "extraction_method": "pymupdf",
                }

            except Exception as e:
                raise TextExtractionError(f"Failed to extract PDF text: {str(e)}")

        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _do_extraction)

    async def _extract_epub(self, content: bytes) -> dict[str, Any]:
        """Extract text from EPUB using ebooklib."""
        def _do_extraction() -> dict[str, Any]:
            text_parts = []
            metadata = {}

            try:
                # Read EPUB from bytes
                book = epub.read_epub(io.BytesIO(content))

                # Extract metadata
                title = book.get_metadata('DC', 'title')
                creator = book.get_metadata('DC', 'creator')
                description = book.get_metadata('DC', 'description')
                language = book.get_metadata('DC', 'language')
                publisher = book.get_metadata('DC', 'publisher')
                date = book.get_metadata('DC', 'date')
                identifier = book.get_metadata('DC', 'identifier')

                metadata = {}
                if title:
                    metadata["title"] = title[0][0] if title else None
                if creator:
                    metadata["author"] = ", ".join([c[0] for c in creator])
                if description:
                    metadata["description"] = description[0][0] if description else None
                if language:
                    metadata["language"] = language[0][0] if language else None
                if publisher:
                    metadata["publisher"] = publisher[0][0] if publisher else None
                if date:
                    metadata["date"] = date[0][0] if date else None
                if identifier:
                    metadata["identifier"] = identifier[0][0] if identifier else None

                # Extract text from all document items
                chapter_count = 0
                for item in book.get_items():
                    if item.get_type() == 9:  # ITEM_DOCUMENT
                        chapter_count += 1
                        content_html = item.get_content().decode('utf-8', errors='ignore')

                        # Parse HTML and extract text
                        soup = BeautifulSoup(content_html, 'html.parser')

                        # Remove script and style elements
                        for script in soup(["script", "style"]):
                            script.decompose()

                        text = soup.get_text(separator='\n', strip=True)
                        if text:
                            text_parts.append(text)

                return {
                    "text": "\n\n".join(text_parts),
                    "page_count": chapter_count,  # Chapters as "pages"
                    "metadata": metadata,
                    "extraction_method": "ebooklib",
                }

            except Exception as e:
                raise TextExtractionError(f"Failed to extract EPUB text: {str(e)}")

        # Run in thread pool
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _do_extraction)

    async def _extract_text_file(self, content: bytes) -> dict[str, Any]:
        """Extract text from plain text file."""
        try:
            # Try common encodings
            text = None
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1251']:
                try:
                    text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                # Fallback: decode with error replacement
                text = content.decode('utf-8', errors='replace')

            # Count lines as "pages"
            line_count = text.count('\n') + 1

            return {
                "text": text,
                "page_count": line_count // 50,  # Approximate pages (50 lines per page)
                "metadata": {},
                "extraction_method": "plain_text",
            }

        except Exception as e:
            raise TextExtractionError(f"Failed to extract text: {str(e)}")

    async def _extract_docx(self, content: bytes) -> dict[str, Any]:
        """Extract text from DOCX using python-docx."""
        def _do_extraction() -> dict[str, Any]:
            text_parts = []
            metadata = {}

            try:
                doc = Document(io.BytesIO(content))

                # Extract metadata from core properties
                core_props = doc.core_properties
                if core_props:
                    metadata = {
                        "title": core_props.title,
                        "author": core_props.author,
                        "subject": core_props.subject,
                        "keywords": core_props.keywords,
                        "category": core_props.category,
                        "comments": core_props.comments,
                        "created": core_props.created.isoformat() if core_props.created else None,
                        "modified": core_props.modified.isoformat() if core_props.modified else None,
                        "last_modified_by": core_props.last_modified_by,
                        "revision": core_props.revision,
                    }
                    # Clean None values
                    metadata = {k: v for k, v in metadata.items() if v}

                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        text_parts.append(text)

                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            text_parts.append(" | ".join(row_text))

                # Estimate page count (approx. 250-300 words per page)
                word_count = sum(len(text.split()) for text in text_parts)
                page_count = max(1, word_count // 275)

                return {
                    "text": "\n\n".join(text_parts),
                    "page_count": page_count,
                    "metadata": metadata,
                    "extraction_method": "python-docx",
                }

            except Exception as e:
                raise TextExtractionError(f"Failed to extract DOCX text: {str(e)}")

        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, _do_extraction)

    async def get_text_preview(
        self,
        content: bytes,
        content_type: str,
        file_name: str,
        max_chars: int = 1000,
    ) -> str:
        """Get a short text preview for display."""
        try:
            result = await self.extract_text(content, content_type, file_name)
            text = result.get("text", "")
            if len(text) > max_chars:
                text = text[:max_chars] + "..."
            return text
        except TextExtractionError:
            return ""


# Singleton instance
_text_extractor: TextExtractor | None = None


def get_text_extractor() -> TextExtractor:
    """Get singleton TextExtractor instance."""
    global _text_extractor
    if _text_extractor is None:
        _text_extractor = TextExtractor()
    return _text_extractor
